import tkinter as tk
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import datetime
from tkcalendar import DateEntry
from tkinter import ttk

def crear_documento():
    # Obtener la fecha seleccionada por el usuario
    fecha = date_picker.get_date()

    # Obtener el nombre de la rama seleccionado por el usuario
    rama_seleccionada = rama_combobox.get()

    # Obtener el nombre de la rama scout ingresado por el usuario
    rama_scout = rama_entry.get()

    # Concatenar los valores de rama seleccionada y rama scout
    rama_completa = f"{rama_seleccionada} - {rama_scout}"

    # Obtener las actividades ingresadas por el usuario
    actividades = []
    for actividad in actividad_entries:
        hora = actividad[0].get()
        nombre = actividad[1].get()
        responsables = actividad[2].get()
        observaciones = actividad[3].get()
        actividades.append((hora, nombre, responsables, observaciones))

    # Obtener la hora de convocación seleccionada por el usuario
    hora_convocacion = f"{horas_combobox.get()}:{minutos_combobox.get()}"
    hora_convocacion_dt = datetime.datetime.strptime(hora_convocacion, "%H:%M")
    concentracion_hora = hora_convocacion
    concentracion_actividad = ("CONCENTRACIÓN", "-", "-")
    actividades.append((concentracion_hora, *concentracion_actividad))
    
    # Crear la actividad de F.I.O 15 minutos después del horario de cita ingresado
    fio_hora = (hora_convocacion_dt + datetime.timedelta(minutes=15)).strftime("%H:%M")
    fio_actividad = ("F.I.O", "Jefe de grupo", "-")
    actividades.append((fio_hora, *fio_actividad))
    
        # Verificar si el checkbox está activado
    if checkbox_var.get():
        # Agregar la actividad de merienda
        merienda_actividad = ("17:00", "MERIENDA", "Equipo de Cocina", "-")
        actividades.append(merienda_actividad)
    
    # Crear un nuevo documento
    doc = Document()

    # Configurar el horario en el encabezado
    header = doc.sections[0].header
    paragraph = header.paragraphs[0]
    paragraph.text = fecha.strftime("%d/%m/%y")  # Formatear la fecha en dd/mm/yy

    # Agregar los párrafos con formato
    doc.add_paragraph("GRUPO SCOUT SAN ANTONIO").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("Centenario 1399 San Antonio de Padua").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("Distrito 5 - Zona 11 BUENOS AIRES NOROESTE").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("Personería Jurídica Nacional: Res. I.G.J. Nº999 del 24/09/1998").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("CUIT 30 – 69732250 – 3 – IVA: Exento").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph().alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Agregar el texto rama_completa
    doc.add_paragraph(rama_completa.upper()).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Crear la tabla
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Configurar los encabezados de la tabla
    headers = table.rows[0].cells
    headers[0].text = 'HORARIO'
    headers[1].text = 'ACTIVIDAD'
    headers[2].text = 'RESPONSABLES'
    headers[3].text = 'OBSERVACIONES'

    # Agregar las actividades a la tabla
    for actividad in actividades:
        hora, nombre, responsables, observaciones = actividad
        
        # Crear una nueva fila en la tabla
        row = table.add_row().cells

        # Agregar los valores a cada celda de la fila
        row[0].text = hora
        row[1].text = nombre
        row[2].text = responsables
        row[3].text = observaciones if observaciones is not None else ''

    # Guardar el documento como un archivo .docx
    direccionDoc = 'Programa_%s.docx' % fecha.strftime("%d-%m-%y")
    doc.save(direccionDoc)
    label.configure(text="¡Documento creado!")

# Crear la ventana de Tkinter
window = tk.Tk()
window.geometry("800x600")
window.title("Programa de Actividades")

# Crear el formulario para ingresar la fecha
date_label = tk.Label(window, text="Fecha:")
date_label.pack()
date_picker = DateEntry(window, width=12, background='darkblue', foreground='white', date_pattern='dd/mm/yyyy')
date_picker.pack()

# Crear el formulario para elegir la rama
rama_label = tk.Label(window, text="Nombre de la Rama:")
rama_label.pack()
rama_frame = tk.Frame(window)
rama_frame.pack()

rama_combobox = ttk.Combobox(rama_frame, values=["COLONIA DE CASTORES", "MANADA", "UNIDAD", "COMUNIDAD CAMINANTES", "COMUNIDAD ROVER"], state="readonly")
rama_combobox.pack(side="left")

rama_entry = tk.Entry(rama_frame)
rama_entry.pack(side="left")


# Crear las listas desplegables para la hora de convocación
hora_convocacion_label = tk.Label(window, text="Hora de Convocación:")
hora_convocacion_label.pack()

horas_combobox = ttk.Combobox(window, values=[str(i).zfill(2) for i in range(24)], state="readonly")
horas_combobox.pack()

minutos_combobox = ttk.Combobox(window, values=[str(i).zfill(2) for i in range(60)], state="readonly")
minutos_combobox.pack()

# Crear el formulario para ingresar las actividades
actividad_entries = []
for i in range(2):
    hora_label = tk.Label(window, text="Hora:")
    hora_label.pack()
    hora_entry = tk.Entry(window)
    hora_entry.pack()

    nombre_label = tk.Label(window, text="Actividad:")
    nombre_label.pack()
    nombre_entry = tk.Entry(window)
    nombre_entry.pack()

    responsables_label = tk.Label(window, text="Responsables:")
    responsables_label.pack()
    responsables_entry = tk.Entry(window)
    responsables_entry.pack()

    observaciones_label = tk.Label(window, text="Observaciones:")
    observaciones_label.pack()
    observaciones_entry = tk.Entry(window)
    observaciones_entry.pack()

    actividad_entries.append((hora_entry, nombre_entry, responsables_entry, observaciones_entry))
    
# Crear un checkbox para agregar una actividad de merienda
checkbox_var = tk.BooleanVar()
checkbox = tk.Checkbutton(window, text="Agregar actividad de merienda", variable=checkbox_var)
checkbox.pack()

# Crear un botón para crear el documento
button_crear = tk.Button(window, text="Crear Documento", command=crear_documento)
button_crear.pack()

# Crear una etiqueta para mostrar el estado
label = tk.Label(window, text="")
label.pack()

# Ejecutar el bucle principal de la ventana
window.mainloop()
